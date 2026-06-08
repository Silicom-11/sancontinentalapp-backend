from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import datetime

desktop_path = r"C:\Users\Equipo\Desktop\Informe_Teoria_General_Sistemas_APA7.docx"

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Title page
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('Aplicación de la Teoría General de Sistemas')
set_font(run, size=20, bold=True)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('\nFacultad de Ingeniería de Sistemas e Informática\nUniversidad Tecnológica del Perú')
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('\nIntegrantes:\n')
set_font(run, size=12, bold=True)
run = p.add_run('Aquino Carhuas Marc Andreessen — U23221709\nRondón Romero Jeancarlos Antony — U23220782\nCastillon Torres Jhonny Alberto — U23226379')
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('\n30 de abril de 2026')
set_font(run, size=12)

# Add page break
from docx.enum.text import WD_BREAK
p.runs[0].add_break(WD_BREAK.PAGE)

# Abstract
h = doc.add_paragraph()
run = h.add_run('Abstract')
set_font(run, size=12, bold=True)

p = doc.add_paragraph()
abstract_text = ("This paper presents an integrative overview of General Systems Theory (GST) and its application across multiple domains: "
                 "open systems, cybernetics and macrosystems, complexity, economics, ecosystems, sociology and politics, organization theory, "
                 "feasible systems, chaos theory, game theory, operations research, system dynamics, classical vs. systems science, systems thinking, "
                 "soft systems methodology, ontological and epistemological approaches, and human systems. Key concepts, applied examples and recommendations "
                 "for research and practice are provided.")
run = p.add_run(abstract_text)
set_font(run, size=12)

# Keywords (Spanish user — include both)
p = doc.add_paragraph()
run = p.add_run('\nKeywords: Teoría General de Sistemas; sistemas abiertos; cibernética; complejidad')
set_font(run, size=12, italic=True)

# Add page break
p.runs[0].add_break(WD_BREAK.PAGE)

# Content sections
sections = [
    ('Introducción', 'La Teoría General de Sistemas proporciona un marco interdisciplinario para comprender entidades compuestas por elementos interrelacionados y sus interacciones con entornos. Este informe sintetiza conceptos y aplicaciones con un enfoque en sistemas reales y problemas socio-tecnológicos.'),
    ('Sistemas abiertos', 'Los sistemas abiertos intercambian materia, energía e información con su entorno. Se analizan principios de homeostasis, equilibrio dinámico, entradas/salidas y retroalimentaciones.'),
    ('Cibernética y macrosistemas', 'La cibernética estudia el control y la comunicación en sistemas. En macrosistemas facilita diseño de mecanismos de regulación y gobernanza.'),
    ('Complejidad', 'Los sistemas complejos exhiben propiedades emergentes, no linealidad y sensibilidad a condiciones iniciales. Se discuten redes, autoorganización y medidas de complejidad.'),
    ('Economía', 'Desde la perspectiva sistémica, la economía se interpreta como un conjunto de agentes interdependientes que intercambian recursos e información.'),
    ('Ecosistemas', 'Los ecosistemas son sistemas abiertos con interacciones tróficas entre especies y el entorno abiótico. Se analiza resiliencia y servicios ecosistémicos.'),
    ('Sociología y política', 'Modelado de influencias, difusión cultural y dinámicas institucionales en sistemas sociotécnicos.'),
    ('Organización', 'La organización vista como sistema enfatiza procesos, estructura, cultura y flujo de información. Se proponen mecanismos para aprendizaje organizacional.'),
    ('Sistemas factibles', 'Análisis de viabilidad técnica, económica y social; evaluación multicriterio.'),
    ('Teoría del caos', 'Sistemas deterministas con comportamiento aparentemente aleatorio; implicaciones para predictibilidad y resiliencia.'),
    ('Teoría de juegos', 'Modela interacciones estratégicas: equilibrios de Nash, cooperación y diseño de mecanismos.'),
    ('Investigación de operaciones', 'Métodos cuantitativos para optimización: programación, colas, redes y simulación.'),
    ('Dinámica de sistemas', 'Modelo de stocks y flujos para representar comportamiento en el tiempo; útil para política y planificación.'),
    ('Ciencia clásica vs ciencia sistémica', 'Comparación entre reduccionismo y enfoque sistémico; complementariedad metodológica.'),
    ('Pensamiento sistémico', 'Enfoque para identificar patrones, relaciones y efectos retardados; evita soluciones parciales.'),
    ('Metodología de sistemas suaves', 'SSM para problemas humanos complejos y mal estructurados: modelos conceptuales y participación.'),
    ('Abordajes ontológicos y epistemológicos', 'Clarificación de supuestos sobre qué existe y cómo se conoce; promueve métodos mixtos.'),
    ('Sistemas humanos', 'Integran emociones, valores y estructuras sociales; requieren enfoques cualitativos y cuantitativos.')
]

for title, text in sections:
    h = doc.add_paragraph()
    run = h.add_run(title)
    set_font(run, size=14, bold=True)
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12)

# Conclusions
h = doc.add_paragraph()
run = h.add_run('Conclusiones y recomendaciones')
set_font(run, size=14, bold=True)

p = doc.add_paragraph()
conclusions = (
    '- Integrar enfoques cuantitativos y cualitativos para mejorar resultados.\n'
    '- Diseñar políticas adaptativas con monitoreo continuo.\n'
    '- Promover la interdisciplinariedad en formación e investigación.\n'
    '- Priorizar la participación de actores en sistemas humanos y sociotécnicos.'
)
run = p.add_run(conclusions)
set_font(run, size=12)

# References (APA 7)
h = doc.add_paragraph()
run = h.add_run('\nReferences')
set_font(run, size=12, bold=True)

refs = [
    "von Bertalanffy, L. (1968). General System Theory: Foundations, Development, Applications. George Braziller.",
    "Ashby, W. R. (1956). An Introduction to Cybernetics. Chapman & Hall.",
    "Forrester, J. W. (1961). Industrial Dynamics. MIT Press.",
    "Meadows, D. H. (2008). Thinking in Systems: A Primer. Chelsea Green Publishing.",
    "Holland, J. H. (1992). Adaptation in Natural and Artificial Systems. MIT Press.",
    "Ostrom, E. (1990). Governing the Commons: The Evolution of Institutions for Collective Action. Cambridge University Press.",
    "Simon, H. A. (1962). The Architecture of Complexity. Proceedings of the American Philosophical Society, 106(6), 467–482.",
    "Nash, J. (1950). Equilibrium points in n-person games. Proceedings of the National Academy of Sciences of the United States of America, 36(1), 48–49."
]

for r in refs:
    p = doc.add_paragraph()
    run = p.add_run(r)
    set_font(run, size=12)

# Save document
try:
    doc.save(desktop_path)
    print('DOCX creado en:', desktop_path)
except Exception as e:
    print('Error al guardar DOCX:', e)

if __name__ == '__main__':
    pass
